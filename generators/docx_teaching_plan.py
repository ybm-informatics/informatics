from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(width_cm * 567)))
    width.set(qn("w:type"), "dxa")


def set_table_grid(table, widths_cm: list[float]) -> None:
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width_cm in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width_cm * 567)))
        grid.append(col)


def set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(int(width_cm * 567)))
    width.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 80, start: int = 80, bottom: int = 80, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_korean_font(run, size: int = 10, bold: bool = False) -> None:
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold


def set_font(paragraph, size: int = 10, bold: bool = False) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    for run in paragraph.runs:
        apply_korean_font(run, size, bold)


def add_cell_text(cell, text: str, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    apply_korean_font(run, size, bold)
    set_font(paragraph, size, bold)


def add_bullets(cell, items: list[str]) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell)
    for index, item in enumerate(items):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        run = paragraph.add_run(f"• {item}")
        apply_korean_font(run)
        set_font(paragraph)


def format_table(table, widths: list[float] | None = None, header_fill: str = "EDEDED") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    if widths:
        set_table_grid(table, widths)
        set_table_width(table, sum(widths))
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)
            if widths:
                set_cell_width(cell, widths[cell_index])
            if row_index == 0:
                set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                set_font(paragraph, 10, row_index == 0)


def load_plan(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def build_docx(plan: dict, output_path: Path, one_page: bool = False) -> None:
    document = Document()
    section = document.sections[0]
    if one_page:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
    else:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("교수·학습 과정안")
    apply_korean_font(title_run, 13 if one_page else 16, True)

    if one_page:
        build_one_page_body(document, plan)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return

    meta = document.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    format_table(meta, [3.0, 13.0])
    meta_data = [
        ("단원", plan["unit_title"]),
        ("소단원", plan["lesson_title"]),
        ("차시", plan["period"]),
        ("교과서", plan["textbook_pages"]),
    ]
    for row, (key, value) in zip(meta.rows, meta_data):
        set_cell_shading(row.cells[0], "F2F2F2")
        add_cell_text(row.cells[0], key, bold=True)
        add_cell_text(row.cells[1], value)

    document.add_paragraph()
    heading = document.add_paragraph()
    heading.add_run("학습 목표").bold = True
    set_font(heading, 11, True)
    for goal in plan["learning_goals"]:
        paragraph = document.add_paragraph(style=None)
        paragraph.add_run(f"• {goal}")
        set_font(paragraph)

    document.add_paragraph()
    method_table = document.add_table(rows=4, cols=2)
    method_table.style = "Table Grid"
    format_table(method_table, [3.0, 13.0])
    rows = [
        ("교수·학습", ", ".join(plan["teaching_learning_methods"])),
        ("평가", ", ".join(plan["assessment_methods"])),
        ("준비물(교사)", ", ".join(plan["materials"]["teacher"])),
        ("준비물(학생)", ", ".join(plan["materials"]["student"])),
    ]
    for row, (key, value) in zip(method_table.rows, rows):
        set_cell_shading(row.cells[0], "F2F2F2")
        add_cell_text(row.cells[0], key, bold=True)
        add_cell_text(row.cells[1], value)

    document.add_paragraph()
    flow_heading = document.add_paragraph()
    flow_heading.add_run("교수·학습 과정").bold = True
    set_font(flow_heading, 11, True)

    flow = document.add_table(rows=1, cols=4)
    flow.style = "Table Grid"
    format_table(flow, [1.6, 3.0, 7.0, 4.4])
    headers = ["단계", "학습 요소", "교수·학습 활동", "지도상의 유의점"]
    for cell, header in zip(flow.rows[0].cells, headers):
        add_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")
    for item in plan["lesson_flow"]:
        row = flow.add_row().cells
        for width, cell in zip([1.6, 3.0, 7.0, 4.4], row):
            set_cell_width(cell, width)
        add_cell_text(row[0], item["stage"], bold=True)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(row[0], "F2F2F2")
        add_cell_text(row[1], "\n".join(item["learning_elements"]))
        add_bullets(row[2], item["teaching_learning_activities"])
        add_bullets(row[3], item["teaching_notes"])

    document.add_paragraph()
    notes_heading = document.add_paragraph()
    notes_heading.add_run("검토 메모").bold = True
    set_font(notes_heading, 11, True)
    for note in plan["review_notes"]:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"• {note}")
        set_font(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def compact_items(items: list[str], max_items: int = 2) -> str:
    return "\n".join(f"• {item}" for item in items[:max_items])


def compact_activities(items: list[str], max_items: int = 3) -> str:
    lines = []
    for item in items[:max_items]:
        if item.startswith(("➊", "➋", "➌", "스스로")):
            lines.append(item)
        elif item.startswith("•"):
            lines.append(item)
        else:
            lines.append(f"• {item}")
    return "\n".join(lines)


def build_one_page_body(document: Document, plan: dict) -> None:
    info = document.add_table(rows=5, cols=4)
    info.style = "Table Grid"
    format_table(info, [1.8, 7.2, 1.8, 6.2])
    info_rows = [
        ("단원", plan["unit_title"], "차시", plan["period"]),
        ("소단원", plan["lesson_title"], "쪽수", plan["textbook_pages"]),
        ("성취기준", ", ".join(plan["achievement_standard_codes"]), "평가", ", ".join(plan["assessment_methods"])),
        ("교수·학습", ", ".join(plan["teaching_learning_methods"]), "준비물", f"교사: {', '.join(plan['materials']['teacher'])}\n학생: {', '.join(plan['materials']['student'])}"),
        ("학습 목표", compact_items(plan["learning_goals"], 2), "", ""),
    ]
    for row, values in zip(info.rows, info_rows):
        if values[0] == "학습 목표":
            merged = row.cells[1].merge(row.cells[3])
            add_cell_text(row.cells[0], values[0], bold=True, size=9)
            set_cell_shading(row.cells[0], "F2F2F2")
            add_cell_text(merged, values[1], size=9)
            set_cell_width(merged, 7.2 + 1.8 + 6.2)
            continue
        for index, value in enumerate(values):
            add_cell_text(row.cells[index], value, bold=index in (0, 2), size=9)
            if index in (0, 2):
                set_cell_shading(row.cells[index], "F2F2F2")

    flow = document.add_table(rows=1, cols=4)
    flow.style = "Table Grid"
    format_table(flow, [1.6, 3.0, 8.0, 4.4])
    headers = ["단계", "학습 요소", "교수·학습 활동", "지도상의 유의점"]
    for cell, header in zip(flow.rows[0].cells, headers):
        add_cell_text(cell, header, bold=True)
        set_cell_shading(cell, "D9EAF7")

    for item in plan["lesson_flow"]:
        row = flow.add_row().cells
        for width, cell in zip([1.6, 3.0, 8.0, 4.4], row):
            set_cell_width(cell, width)
        add_cell_text(row[0], item["stage"], bold=True)
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(row[0], "F2F2F2")
        add_cell_text(row[1], "\n".join(item["learning_elements"][:4]), size=9)
        max_activities = 12 if item["stage"] == "전개" else 3
        add_cell_text(row[2], compact_activities(item["teaching_learning_activities"], max_activities), size=8)
        add_cell_text(row[3], compact_activities(item["teaching_notes"], 3), size=8)

    note_table = document.add_table(rows=1, cols=2)
    note_table.style = "Table Grid"
    format_table(note_table, [2.2, 14.8])
    add_cell_text(note_table.rows[0].cells[0], "비고", bold=True, size=8)
    set_cell_shading(note_table.rows[0].cells[0], "F2F2F2")
    add_cell_text(note_table.rows[0].cells[1], "차시 시간, 활동 시간, 평가 문항은 학교 운영 계획에 맞게 조정한다.", size=8)


def main() -> int:
    parser = argparse.ArgumentParser(description="Create DOCX from teaching plan draft JSON.")
    parser.add_argument(
        "--input",
        type=Path,
        default=Path("outputs/teaching_plans/unit_3_01_1_teaching_plan_draft.json"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("outputs/teaching_plans/unit_3_01_1_teaching_plan_draft.docx"),
    )
    parser.add_argument("--one-page", action="store_true")
    args = parser.parse_args()
    plan = load_plan(args.input)
    build_docx(plan, args.output, one_page=args.one_page)
    print("DOCX 교수학습과정안을 생성했습니다.")
    print(f"- {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
